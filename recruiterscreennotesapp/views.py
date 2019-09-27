from django.shortcuts import render

from docx import Document

# Create your views here.


  
from django.shortcuts import render

# Create your views here.
# recruiterscreen



from .forms import RecruiterScreenForm
# from .models import RecruiterScreen

# Create your views here.


def recruiterscreen(request):
    document = Document()

    title = 'Recruiter Screen Notes'
    form = RecruiterScreenForm()
    if request.method == "POST":
        form = RecruiterScreenForm(request.POST or None)
        if form.is_valid():
            print(form.cleaned_data)
            candidatename = form.cleaned_data['candidatename']
            email = form.cleaned_data['email']
            screen_date = form.cleaned_data['screen_date']
            experience = form.cleaned_data['experience']
            leadership = form.cleaned_data['leadership']
            motivation = form.cleaned_data['motivation']
            compensation = form.cleaned_data['compensation']
            visa1 = form.cleaned_data['visa1']
            visa2 = form.cleaned_data['visa2']
            visa3 = form.cleaned_data['visa3']
            visa4 = form.cleaned_data['visa4']
            visa5 = form.cleaned_data['visa5']
            visa6 = form.cleaned_data['visa6']
            visa7 = form.cleaned_data['visa7']
            additional_comments = form.cleaned_data['additional_comments']

            def makefilenameandpath():
                new_filename = ('Recruiter Screen - ' + candidatename + '-' + str(screen_date))
                ext = 'docx'
                final_filename = '{new_filename}.{ext}'.format(new_filename=new_filename, ext=ext)
                print(final_filename)
                return '{final_filename}'.format(new_filename=new_filename, final_filename=final_filename)

            docname = makefilenameandpath()

            document.add_heading('Recruiter Screen - ' + candidatename)

            document.add_heading('Candidate Details', level=1)
            p1=document.add_paragraph('Date: ')
            p1.add_run(str(screen_date))
            p2=document.add_paragraph('Name: ')
            p2.add_run(candidatename)
            p3=document.add_paragraph('Email: ')
            p3.add_run(email)

            document.add_heading('Screen Information', level=1)
            p4=document.add_paragraph('Experience: ')
            p4.add_run(experience)
            p5 = document.add_paragraph('Leadership: ')
            p5.add_run(leadership)
            p6 = document.add_paragraph('Motivation: ')
            p6.add_run(motivation)
            p7 = document.add_paragraph('Compensation: ')
            p7.add_run(compensation)
            p8= document.add_paragraph('Additional Comments: ')
            p8.add_run(additional_comments)

            document.add_heading('Visa Information', level=1)
            p9=document.add_paragraph('''What type of visa status do they currently hold?''')
            p9.add_run(visa2)
            p9.add_run(' ')
            p9.add_run(visa7)
            p10=document.add_paragraph('''Does the candidate have the unrestricted right to work indefinitely in the United States? What is the candidate's current status?: ''')
            p10.add_run(visa1)
            p11 =document.add_paragraph('''If sponsorship is required, when does the candidate's current U.S. Immigration status expire?: ''')
            p11.add_run(visa3)
            p12 = document.add_paragraph('''If in H-1B status, how much time does he or she have remaining towards the 6 year limit?''')
            p12.add_run(visa4)
            p13 = document.add_paragraph('''If on OPT, is the candidate eligible for a STEM extension? If so, what is the expected expiration date of their STEM extension?''')
            p13.add_run(visa5)
            p14 =document.add_paragraph('''Is the candidate in the green card process? Does he or she have a certified PERM? Does he or she have an approved I-140, and have 180 days passed since approval?''')
            p14.add_run(visa6)


            document.save(docname)

            form = RecruiterScreenForm()

            


    context = {
        'title' : title,
        'form' : form,


    }
    return render(request, "recruiterscreennotesapp/recruiternotes.html", context)
